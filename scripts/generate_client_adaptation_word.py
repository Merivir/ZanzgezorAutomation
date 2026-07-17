import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "CLIENT_ADAPTATION_GUIDE.md"
OUTPUT = ROOT / "docs" / "Client_Adaptation_Guide.docx"

NAVY = "17324D"
TEAL = "197C82"
PALE_TEAL = "E8F3F3"
PALE_BLUE = "EEF3F8"
MID_GREY = "66717D"
LIGHT_GREY = "E2E7EB"
CODE_BG = "F4F6F8"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Client Adaptation Guide  |  ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(begin)
    paragraph._p.append(instruction)
    paragraph._p.append(end)


def add_inline_text(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Cascadia Mono"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(NAVY)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("263746")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in (
        ("Title", 30, NAVY),
        ("Subtitle", 13, MID_GREY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 13, TEAL),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)

    code = styles.add_style("Code Block", 1)
    code.font.name = "Cascadia Mono"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string("203040")
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.16)
    code.paragraph_format.space_before = Pt(5)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.keep_together = True


def add_cover(document):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(70)

    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("AUTOMATION FRAMEWORK")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Client Adaptation Guide")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("A practical handover for configuring, validating, and extending the QA automation project")

    line = document.add_table(rows=1, cols=1)
    line.alignment = WD_TABLE_ALIGNMENT.CENTER
    line.autofit = False
    line.columns[0].width = Inches(2.2)
    cell = line.cell(0, 0)
    set_cell_shading(cell, TEAL)
    cell.height = Pt(3)
    cell.text = ""

    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_before = Pt(28)
    run = summary.add_run("Includes configuration, architecture, SIP adaptation, test catalogue migration, execution, logging, and acceptance criteria.")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)

    catalogue = document.add_paragraph()
    catalogue.alignment = WD_ALIGN_PARAGRAPH.CENTER
    catalogue.paragraph_format.space_before = Pt(16)
    run = catalogue.add_run("Reference catalogue: 120 call scenarios across 5 functional suites")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_before = Pt(95)
    run = date.add_run("Version 1.0  |  July 2026")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)

    document.add_page_break()


def parse_table(lines):
    rows = []
    for line in lines:
        values = [value.strip() for value in line.strip().strip("|").split("|")]
        rows.append(values)
    return [rows[0]] + rows[2:]


def add_table(document, rows):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, PALE_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_text(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(document, lines):
    paragraph = document.add_paragraph(style="Code Block")
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_BG)
    properties.append(shading)
    paragraph.add_run("\n".join(lines))


def add_body(document, markdown):
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines = []
    skipped_title = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if stripped.startswith("# ") and not skipped_title:
            skipped_title = True
            index += 1
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
            index += 1
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\|\s*-", lines[index + 1]):
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, parse_table(table_lines))
            continue

        if re.match(r"^-\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_text(paragraph, re.sub(r"^-\s+", "", stripped))
            index += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            add_inline_text(paragraph, re.sub(r"^\d+\.\s+", "", stripped))
            index += 1
            continue

        if stripped:
            paragraph = document.add_paragraph()
            add_inline_text(paragraph, stripped)

        index += 1


def build_document():
    document = Document()
    configure_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    add_cover(document)
    add_body(document, SOURCE.read_text(encoding="utf-8"))

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        add_page_number(footer)
        for run in footer.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MID_GREY)

    properties = document.core_properties
    properties.title = "Client Adaptation Guide"
    properties.subject = "QA automation framework onboarding and client adaptation"
    properties.author = "QA Automation Team"
    properties.keywords = "QA, automation, Selenium, SIP, PJSUA, client onboarding"

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
